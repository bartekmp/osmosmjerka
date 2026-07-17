from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from osmosmjerka.logging_config import get_logger
from PIL import Image, ImageDraw, ImageFont

logger = get_logger(__name__)


def _shade_cell(cell_obj, fill_hex: str, pattern: str = "clear", pattern_color: str = "auto") -> None:
    """Set a table cell's background fill/pattern (used for crossword blank/blocked cells)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), pattern)
    shd.set(qn("w:color"), pattern_color)
    shd.set(qn("w:fill"), fill_hex)
    cell_obj._tc.get_or_add_tcPr().append(shd)


def _crop_crossword_grid(grid: list) -> tuple[list, int, int]:
    """Trim a crossword grid down to its used content plus a 1-cell border.

    The generator always allocates a grid sized for the longest word (at least
    10x10), so a puzzle using only a handful of short phrases can leave most of
    the grid empty. Exporting that full allocation makes for a mostly-blank,
    oversized image/table, so this crops to the bounding box of placed letters.

    Returns (cropped_grid, row_offset, col_offset) — the offsets to add to a
    cropped cell's (row, col) to recover its coordinates in the original grid,
    needed to look up clue-start numbers (keyed by original coordinates).
    """
    rows_with_content = [r for r, row in enumerate(grid) if any(cell is not None for cell in row)]
    if not rows_with_content:
        return grid, 0, 0
    cols_with_content = [c for row in grid for c, cell in enumerate(row) if cell is not None]

    pad = 1
    min_r = max(0, min(rows_with_content) - pad)
    max_r = min(len(grid) - 1, max(rows_with_content) + pad)
    min_c = max(0, min(cols_with_content) - pad)
    max_c = min(len(grid[0]) - 1, max(cols_with_content) + pad)

    cropped = [row[min_c : max_c + 1] for row in grid[min_r : max_r + 1]]
    return cropped, min_r, min_c


def _crossword_start_numbers(phrases: list) -> dict:
    """Map (row, col) -> clue number label for each phrase's starting cell.

    When multiple phrases start at the same cell, their numbers are joined with
    "/" (e.g. "1/2"), matching the live crossword grid's numbering.
    """
    numbers: dict[tuple[int, int], list] = {}
    for p in phrases:
        coords = p.get("coords")
        number = p.get("start_number")
        if coords and number:
            r, c = coords[0]
            numbers.setdefault((r, c), []).append(number)
    return {key: "/".join(str(n) for n in nums) for key, nums in numbers.items()}


def _crossword_clues(phrases: list) -> tuple[list[str], list[str]]:
    """Build numbered "Across"/"Down" clue lines from placed crossword phrases.

    The clue text is the translation (the puzzle-solver never sees the answer word
    itself), matching how clues are shown in the live crossword UI.
    """

    def clue_text(p: dict) -> str:
        return f"{p.get('start_number')}. {p.get('translation') or p.get('phrase', '')}"

    across = sorted((p for p in phrases if p.get("direction") == "across"), key=lambda p: p.get("start_number") or 0)
    down = sorted((p for p in phrases if p.get("direction") == "down"), key=lambda p: p.get("start_number") or 0)
    return [clue_text(p) for p in across], [clue_text(p) for p in down]


def _add_docx_clue_section(doc: Document, label: str, lines: list[str]) -> None:
    if not lines:
        return
    heading = doc.add_paragraph()
    heading_run = heading.add_run(label.upper())
    heading_run.bold = True
    heading_run.font.size = Pt(13)
    for line in lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)


def export_to_docx(
    category: str,
    grid: list,
    phrases: list,
    game_type: str = "word_search",
    across_label: str = "Across",
    down_label: str = "Down",
) -> bytes:
    """Export the puzzle grid and phrases to a DOCX file.

    Args:
        category (str): The category of the puzzle.
        grid (list): The puzzle grid as a list of lists. Cells may be `None` for
            crossword blank/blocked squares, which are rendered as shaded cells.
        phrases (list): A list of dictionaries with "phrase" and "translation" keys.
            Crossword phrases additionally carry "coords", "direction", "start_number".
        game_type (str): "word_search" or "crossword". Word search exports the
            fully solved grid (all letters shown, as in normal play). Crossword
            exports a blank, printable puzzle: numbered cells with no letters,
            paired with a numbered Across/Down clue list.
        across_label (str): Localized "Across" heading, used only for crosswords.
        down_label (str): Localized "Down" heading, used only for crosswords.

    Returns:
        bytes: The DOCX file content as bytes.
    """
    if not category or not grid or not phrases:
        raise ValueError("Category, grid, and phrases must be provided.")
    is_crossword = game_type == "crossword"
    row_offset = col_offset = 0
    if is_crossword:
        grid, row_offset, col_offset = _crop_crossword_grid(grid)
    doc = Document()
    # Uppercase header with category name, centered
    heading = doc.add_heading(category.upper(), 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Letter grid, monospace font, centered, no borders
    table = doc.add_table(rows=len(grid), cols=len(grid[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False  # Turn off auto-fit

    # Remove all borders
    tbl = table._tbl
    for border in tbl.xpath(".//w:tblBorders"):
        border.getparent().remove(border)

    cell_dim = Cm(1.0) if is_crossword else Cm(0.8)
    numbers = _crossword_start_numbers(phrases) if is_crossword else {}

    for r, row in enumerate(grid):
        for c, cell in enumerate(row):
            cell_obj = table.cell(r, c)
            cell_obj.width = cell_dim
            if cell is None:
                if is_crossword:
                    # Gray diagonal-stripe pattern for non-playable squares.
                    _shade_cell(cell_obj, "D9D9D9", pattern="diagStripe", pattern_color="808080")
                else:
                    _shade_cell(cell_obj, "000000")
                continue
            p = cell_obj.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            number = numbers.get((r + row_offset, c + col_offset))
            if number:
                num_run = p.add_run(number)
                num_run.font.superscript = True
                num_run.font.name = "Arial"
                num_run.font.size = Pt(6)
            if is_crossword:
                # A crossword is exported blank, for solving on paper — the answer
                # letters aren't shown, only the clue numbers.
                continue
            letter_run = p.add_run(cell)
            letter_run.font.name = "Courier New"
            letter_run.font.size = Pt(12)
            letter_run.bold = True  # Make grid letters bold

    # Set row height for uniform appearance
    for row in table.rows:
        row.height = cell_dim

    doc.add_paragraph()  # Add a blank line for spacing

    if is_crossword:
        across_lines, down_lines = _crossword_clues(phrases)
        _add_docx_clue_section(doc, across_label, across_lines)
        _add_docx_clue_section(doc, down_label, down_lines)
    else:
        # Convert phrases to uppercase and join them with commas
        phrases_line = ", ".join(w["phrase"].upper() for w in phrases)
        p = doc.add_paragraph(phrases_line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.bold = True  # Make phrases bold

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _load_export_fonts(cell_size: int) -> dict:
    try:
        return {
            "title": ImageFont.truetype("fonts/DejaVuSans-Bold.ttf", 24),
            "grid": ImageFont.truetype("fonts/DejaVuSans-Bold.ttf", max(18, cell_size // 2)),
            "phrases": ImageFont.truetype("fonts/DejaVuSans-Bold.ttf", 16),
            "header": ImageFont.truetype("fonts/DejaVuSans-Bold.ttf", 18),
            "number": ImageFont.truetype("fonts/DejaVuSans-Bold.ttf", max(10, cell_size // 3)),
        }
    except OSError:
        try:
            return {
                "title": ImageFont.truetype("fonts/arialbd.ttf", 24),
                "grid": ImageFont.truetype("fonts/arialbd.ttf", max(18, cell_size // 2)),
                "phrases": ImageFont.truetype("fonts/arialbd.ttf", 16),
                "header": ImageFont.truetype("fonts/arialbd.ttf", 18),
                "number": ImageFont.truetype("fonts/arialbd.ttf", max(10, cell_size // 3)),
            }
        except OSError:
            logger.warning("Could not load custom fonts, falling back to default font")
            default = ImageFont.load_default()
            return dict.fromkeys(("title", "grid", "phrases", "header", "number"), default)


def _wrap_text_line(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split(" ")
    lines: list[str] = []
    current = ""
    for word in words:
        test = f"{current} {word}".strip()
        bbox = draw.textbbox((0, 0), test, font=font)
        if bbox[2] - bbox[0] <= max_width or not current:
            current = test
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def _draw_blank_cell(draw: ImageDraw.ImageDraw, x: int, y: int, cell_size: int) -> None:
    """Gray diagonal-stripe pattern for a crossword's non-playable squares."""
    draw.rectangle([x, y, x + cell_size, y + cell_size], fill="#d9d9d9", outline="black", width=1)
    spacing = 6
    for offset in range(-cell_size, cell_size, spacing):
        x0, y0 = x + max(0, offset), y + max(0, -offset)
        x1, y1 = x + min(cell_size, offset + cell_size), y + min(cell_size, -offset + cell_size)
        if x0 < x1 and y0 < y1:
            draw.line([x0, y0, x1, y1], fill="#888888", width=1)


def export_to_png(
    category: str,
    grid: list,
    phrases: list,
    game_type: str = "word_search",
    across_label: str = "Across",
    down_label: str = "Down",
) -> bytes:
    """Export the puzzle grid and phrases to a PNG image.

    Args:
        category (str): The category of the puzzle.
        grid (list): The puzzle grid as a list of lists. Cells may be `None` for
            crossword blank/blocked squares, rendered as a gray diagonal-stripe pattern.
        phrases (list): A list of dictionaries with "phrase" and "translation" keys.
            Crossword phrases additionally carry "coords", "direction", "start_number".
        game_type (str): "word_search" or "crossword". Word search exports the
            fully solved grid (all letters shown, as in normal play). Crossword
            exports a blank, printable puzzle: numbered cells with no letters,
            paired with a numbered Across/Down clue list.
        across_label (str): Localized "Across" heading, used only for crosswords.
        down_label (str): Localized "Down" heading, used only for crosswords.

    Returns:
        bytes: The PNG file content as bytes.
    """
    if not category or not grid or not phrases:
        raise ValueError("Category, grid, and phrases must be provided.")

    is_crossword = game_type == "crossword"
    row_offset = col_offset = 0
    if is_crossword:
        grid, row_offset, col_offset = _crop_crossword_grid(grid)

    grid_rows = len(grid)
    grid_cols = len(grid[0])
    cell_size = max(30, min(50, 800 // max(grid_rows, grid_cols)))  # Adaptive cell size
    margin = 50
    title_height = 60
    grid_width = grid_cols * cell_size
    grid_height = grid_rows * cell_size
    image_width = max(grid_width + (2 * margin), 700) if is_crossword else grid_width + (2 * margin)

    fonts = _load_export_fonts(cell_size)
    numbers = _crossword_start_numbers(phrases) if is_crossword else {}

    # Crossword clue lines (numbered, split Across/Down) need to be measured and
    # word-wrapped up front so the image can be sized to fit them, unlike word
    # search's short comma-joined phrase list which fits a fixed-height footer.
    measurer = ImageDraw.Draw(Image.new("RGB", (1, 1)))
    line_height = 22
    render_lines: list[tuple[str, str]] = []  # (text, "header" | "clue")
    if is_crossword:
        max_clue_width = image_width - (2 * margin)
        across_lines, down_lines = _crossword_clues(phrases)
        for label, lines in ((across_label, across_lines), (down_label, down_lines)):
            if not lines:
                continue
            render_lines.append((label.upper(), "header"))
            for line in lines:
                for wrapped in _wrap_text_line(measurer, line, fonts["phrases"], max_clue_width):
                    render_lines.append((wrapped, "clue"))
        phrases_height = 30 + len(render_lines) * line_height + 8
    else:
        phrases_height = 100

    image_height = grid_height + title_height + phrases_height + (2 * margin)
    img = Image.new("RGB", (image_width, image_height), "white")
    draw = ImageDraw.Draw(img)

    # Draw title
    title_text = category.upper()
    title_bbox = draw.textbbox((0, 0), title_text, font=fonts["title"])
    title_width = title_bbox[2] - title_bbox[0]
    title_x = (image_width - title_width) // 2
    draw.text((title_x, margin), title_text, fill="black", font=fonts["title"])

    # Draw grid
    grid_start_y = margin + title_height
    grid_start_x = (image_width - grid_width) // 2

    for i, row in enumerate(grid):
        for j, cell in enumerate(row):
            x = grid_start_x + (j * cell_size)
            y = grid_start_y + (i * cell_size)

            if cell is None:
                if is_crossword:
                    _draw_blank_cell(draw, x, y, cell_size)
                else:
                    draw.rectangle([x, y, x + cell_size, y + cell_size], fill="black")
                continue

            # Draw cell border
            draw.rectangle([x, y, x + cell_size, y + cell_size], outline="black", width=1)

            if is_crossword:
                # A crossword is exported blank, for solving on paper — the answer
                # letters aren't shown, only the clue numbers.
                number = numbers.get((i + row_offset, j + col_offset))
                if number:
                    draw.text((x + 2, y + 1), number, fill="black", font=fonts["number"])
                continue

            # Draw cell text
            text_bbox = draw.textbbox((0, 0), cell, font=fonts["grid"])
            text_width = text_bbox[2] - text_bbox[0]
            text_height = text_bbox[3] - text_bbox[1]
            text_x = x + (cell_size - text_width) // 2
            text_y = y + (cell_size - text_height) // 2
            draw.text((text_x, text_y), cell, fill="black", font=fonts["grid"])

    phrases_y = grid_start_y + grid_height + 30

    if is_crossword:
        y_cursor = phrases_y
        first = True
        for text, kind in render_lines:
            if kind == "header":
                if not first:
                    y_cursor += 8
                draw.text((margin, y_cursor), text, fill="black", font=fonts["header"])
            else:
                draw.text((margin, y_cursor), text, fill="black", font=fonts["phrases"])
            y_cursor += line_height
            first = False
    else:
        # Draw phrases, centered, wrapped into multiple comma-joined lines
        phrases_text = ", ".join(w["phrase"].upper() for w in phrases)
        max_width = image_width - (2 * margin)
        phrases_lines = []
        current_line = ""

        for phrase in phrases_text.split(", "):
            test_line = current_line + (", " if current_line else "") + phrase
            test_bbox = draw.textbbox((0, 0), test_line, font=fonts["phrases"])
            test_width = test_bbox[2] - test_bbox[0]

            if test_width <= max_width:
                current_line = test_line
            else:
                if current_line:
                    phrases_lines.append(current_line)
                current_line = phrase

        if current_line:
            phrases_lines.append(current_line)

        for i, line in enumerate(phrases_lines):
            line_bbox = draw.textbbox((0, 0), line, font=fonts["phrases"])
            line_width = line_bbox[2] - line_bbox[0]
            line_x = (image_width - line_width) // 2
            line_y = phrases_y + (i * 20)
            draw.text((line_x, line_y), line, fill="black", font=fonts["phrases"])

    # Save to BytesIO
    buffer = BytesIO()
    img.save(buffer, format="PNG")
    return buffer.getvalue()
