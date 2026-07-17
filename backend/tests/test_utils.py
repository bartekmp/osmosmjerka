import pytest
from osmosmjerka.utils import export_to_docx


def test_export_to_docx_valid_input(tmp_path):
    category = "Animals"
    grid = [["C", "A", "T"], ["D", "O", "G"], ["F", "O", "X"]]
    phrases = [
        {"phrase": "cat", "translation": "kot"},
        {"phrase": "dog", "translation": "pies"},
        {"phrase": "fox", "translation": "lis"},
    ]
    docx_bytes = export_to_docx(category, grid, phrases)
    assert isinstance(docx_bytes, bytes)
    # Save to file and check it's a valid docx
    docx_file = tmp_path / "test.docx"
    docx_file.write_bytes(docx_bytes)
    # Try to open with python-docx
    from docx import Document

    doc = Document(str(docx_file))
    assert doc.paragraphs[0].text == category.upper()
    assert phrases[0]["phrase"].upper() in doc.paragraphs[-1].text


def test_export_to_docx_empty_category():
    grid = [["A"]]
    phrases = [{"phrase": "a", "translation": "A"}]
    with pytest.raises(ValueError):
        export_to_docx("", grid, phrases)


def test_export_to_docx_empty_grid():
    category = "Test"
    phrases = [{"phrase": "a", "translation": "A"}]
    with pytest.raises(ValueError):
        export_to_docx(category, [], phrases)


def test_export_to_docx_empty_phrases():
    category = "Test"
    grid = [["A"]]
    with pytest.raises(ValueError):
        export_to_docx(category, grid, [])


def test_export_to_docx_single_letter_grid(tmp_path):
    category = "Single"
    grid = [["Z"]]
    phrases = [{"phrase": "z", "translation": "Z"}]
    docx_bytes = export_to_docx(category, grid, phrases)
    assert isinstance(docx_bytes, bytes)
    docx_file = tmp_path / "single.docx"
    docx_file.write_bytes(docx_bytes)
    from docx import Document

    doc = Document(str(docx_file))
    assert doc.tables[0].cell(0, 0).text == "Z"


def test_export_to_docx_phrases_with_special_characters(tmp_path):
    category = "Special"
    grid = [["Ä", "Ö"], ["Ü", "ß"]]
    phrases = [{"phrase": "äöüß", "translation": "special"}]
    docx_bytes = export_to_docx(category, grid, phrases)
    docx_file = tmp_path / "special.docx"
    docx_file.write_bytes(docx_bytes)
    from docx import Document

    doc = Document(str(docx_file))
    assert "Ä" in doc.tables[0].cell(0, 0).text
    assert "Ö" in doc.tables[0].cell(0, 1).text
    assert "Ü" in doc.tables[0].cell(1, 0).text
    assert "ß" in doc.tables[0].cell(1, 1).text


def test_export_to_png_valid_input(tmp_path):
    from osmosmjerka.utils import export_to_png

    category = "Animals"
    grid = [["C", "A", "T"], ["D", "O", "G"], ["F", "O", "X"]]
    phrases = [
        {"phrase": "cat", "translation": "kot"},
        {"phrase": "dog", "translation": "pies"},
        {"phrase": "fox", "translation": "lis"},
    ]
    png_bytes = export_to_png(category, grid, phrases)
    assert isinstance(png_bytes, bytes)
    png_file = tmp_path / "test.png"
    png_file.write_bytes(png_bytes)
    from PIL import Image

    img = Image.open(str(png_file))
    assert img.format == "PNG"


def test_export_to_png_empty_phrases():
    from osmosmjerka.utils import export_to_png

    category = "Test"
    grid = [["A"]]
    with pytest.raises(ValueError):
        export_to_png(category, grid, [])


CROSSWORD_GRID = [["C", "A", "T"], [None, None, "O"], [None, None, "X"]]
CROSSWORD_PHRASES = [
    {"phrase": "cat", "translation": "kot", "coords": [[0, 0], [0, 1], [0, 2]], "direction": "across", "start_number": 1},
    {"phrase": "tox", "translation": "tox", "coords": [[0, 2], [1, 2], [2, 2]], "direction": "down", "start_number": 2},
]


def test_export_to_docx_crossword_crops_oversized_allocation(tmp_path):
    """A puzzle placed in a small corner of an oversized generator allocation should
    export a small table, not a mostly-blank one matching the full allocation."""
    size = 10
    grid = [[None for _ in range(size)] for _ in range(size)]
    grid[4][4], grid[4][5], grid[4][6] = "C", "A", "T"
    phrases = [{"phrase": "cat", "translation": "kot", "coords": [[4, 4], [4, 5], [4, 6]], "direction": "across", "start_number": 1}]

    docx_bytes = export_to_docx("Small", grid, phrases, game_type="crossword")
    docx_file = tmp_path / "small_crossword.docx"
    docx_file.write_bytes(docx_bytes)
    from docx import Document

    doc = Document(str(docx_file))
    table = doc.tables[0]
    assert len(table.rows) < size
    assert len(table.columns) < size
    # The clue number should still land on the right cell after cropping.
    assert table.cell(1, 1).text == "1"


def test_crop_crossword_grid_trims_unused_allocation():
    """The generator always allocates a grid sized for the longest word (>=10x10),
    so a puzzle using only a few short phrases leaves most of it empty. The export
    should crop to the used content (plus a 1-cell border), not the full allocation."""
    from osmosmjerka.utils import _crop_crossword_grid

    size = 10
    grid = [[None for _ in range(size)] for _ in range(size)]
    # A tiny 3-letter word placed near the middle of the oversized allocation.
    grid[4][4], grid[4][5], grid[4][6] = "C", "A", "T"

    cropped, row_offset, col_offset = _crop_crossword_grid(grid)
    assert len(cropped) < size
    assert len(cropped[0]) < size
    assert row_offset == 3  # one cell of padding above row 4
    assert col_offset == 3  # one cell of padding left of col 4
    # The letters should still be reachable at their offset-adjusted position.
    assert cropped[4 - row_offset][4 - col_offset] == "C"


def test_crossword_start_numbers_shared_cell():
    """When two phrases start at the same cell, their numbers are joined with "/",
    matching the live crossword grid's numbering convention."""
    from osmosmjerka.utils import _crossword_start_numbers

    phrases = [
        {"coords": [[0, 0], [0, 1]], "direction": "across", "start_number": 1},
        {"coords": [[0, 0], [1, 0]], "direction": "down", "start_number": 2},
        {"coords": [[2, 2], [2, 3]], "direction": "across", "start_number": 3},
    ]
    numbers = _crossword_start_numbers(phrases)
    assert numbers[(0, 0)] == "1/2"
    assert numbers[(2, 2)] == "3"


def test_export_to_docx_crossword_is_blank_with_clues(tmp_path):
    """Crossword export is a blank, printable puzzle: numbered cells with no answer
    letters, non-playable cells shaded, and a numbered Across/Down clue list (using
    the translation, not the answer) instead of a plain word list."""
    docx_bytes = export_to_docx("Crossword", CROSSWORD_GRID, CROSSWORD_PHRASES, game_type="crossword")
    assert isinstance(docx_bytes, bytes)
    docx_file = tmp_path / "crossword.docx"
    docx_file.write_bytes(docx_bytes)
    from docx import Document

    doc = Document(str(docx_file))
    assert doc.tables[0].cell(0, 0).text == "1"  # clue number only, no letter
    assert doc.tables[0].cell(0, 1).text == ""  # playable but not a clue start: fully blank
    assert doc.tables[0].cell(1, 0).text == ""  # blank/shaded cell, no text
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "ACROSS" in body_text
    assert "DOWN" in body_text
    assert "1. kot" in body_text
    assert "2. tox" in body_text
    # The answer words themselves should not leak into the clue text or the grid
    assert "1. cat" not in body_text
    assert "cat" not in body_text.lower()


def test_export_to_png_crossword_blank_cells(tmp_path):
    """Crossword export is a blank puzzle: non-playable cells get a gray stripe
    pattern, playable cells get numbers but no letters, plus an Across/Down clue list."""
    from osmosmjerka.utils import export_to_png

    png_bytes = export_to_png("Crossword", CROSSWORD_GRID, CROSSWORD_PHRASES, game_type="crossword")
    assert isinstance(png_bytes, bytes)
    png_file = tmp_path / "crossword.png"
    png_file.write_bytes(png_bytes)
    from PIL import Image

    img = Image.open(str(png_file))
    assert img.format == "PNG"
