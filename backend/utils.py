from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def export_to_docx(category: str, grid: list, words: list) -> bytes:
    """Export the word search grid and words to a DOCX file.

    Args:
        category (str): The category of the word search.
        grid (list): The word search grid as a list of lists.
        words (list): A list of dictionaries with "word" and "translation" keys.

    Returns:
        bytes: The DOCX file content as bytes.
    """
    if not category or not grid or not words:
        raise ValueError("Category, grid, and words must be provided.")
    doc = Document()
    # Uppercase header with category name, centered
    heading = doc.add_heading(category.upper(), 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Letter grid, monospace font, centered, no borders
    table = doc.add_table(rows=len(grid), cols=len(grid[0]))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False  # Turn off auto-fit

    # Remove all borders
    tbl = table._tbl
    for border in tbl.xpath(".//w:tblBorders"):
        border.getparent().remove(border)

    # Set the cell width and height to 0.8 cm
    for r, row in enumerate(grid):
        for c, cell in enumerate(row):
            cell_obj = table.cell(r, c)
            cell_obj.text = cell
            cell_obj.width = Cm(0.8)
            cell_obj.height = Cm(0.8)
            for p in cell_obj.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(12)

    doc.add_paragraph()  # Add a blank line for spacing

    # Convert words to uppercase and join them with commas
    words_line = ", ".join(w["word"].upper() for w in words)
    p = doc.add_paragraph(words_line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()
